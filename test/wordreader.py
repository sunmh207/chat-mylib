import docx

def parse(file_path):
    # with open(file_path, 'r', encoding="ISO-8859-1") as f:
    #     content = f.read()
    #     print(content)

    # exit(0)

    # 打开Word文件
    doc = docx.Document(file_path)
    # 读取每一段的内容
    text = ''
    for para in doc.paragraphs:
        print(para.text)
        print('=========')
        text += para.text

    # 返回解析后的文本内容
    return text
def parse_bypage(file_path):
    from docx import Document
    fn = '1.doc'
    document = Document(file_path)
    pn = 1
    import re
    for p in document.paragraphs:
        r = re.match('Chapter \d+', p.text)
        if r:
            print(r.group(), pn)
        for run in p.runs:
            if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                pn += 1
                print('!!', '=' * 50, pn)

if __name__ == '__main__':
    import sys
    file_path = sys.argv[1]
    print(file_path)

    parse_bypage(file_path)